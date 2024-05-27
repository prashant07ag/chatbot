import os
from io import BytesIO
from dotenv import load_dotenv
from docx import Document
import streamlit as st
import pyperclip
from streamlit_chat import message
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferWindowMemory
from langchain.prompts import (
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate,
    ChatPromptTemplate,
    MessagesPlaceholder
)
from langchain_google_genai.chat_models import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from redundant_filter_retriever import RedundantFilterRetriever
import asyncio

# Helper function to get or create an event loop
def get_or_create_eventloop():
    try:
        return asyncio.get_event_loop()
    except RuntimeError as ex:
        if "There is no current event loop in thread" in str(ex):
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            return asyncio.get_event_loop()

# Initialize event loop
loop = asyncio.new_event_loop()
asyncio.set_event_loop(loop)

# Load environment variables
load_dotenv()

# Set Streamlit page configuration
st.set_page_config(
    page_title="MeitY Guidelines BOT",
    page_icon="🤖",
    layout="wide"
)

# Sidebar setup
st.sidebar.title("Ministry of Electronics and Information Technology")
st.sidebar.markdown("Welcome to MietY Guidelines QnA!")
st.sidebar.markdown("Here are some useful links:")
st.sidebar.markdown("[MeitY Homepage](https://www.meity.gov.in)")
st.sidebar.markdown("[Guidelines](https://www.meity.gov.in/policies-guidelines)")
st.sidebar.markdown("[FAQs](https://www.meity.gov.in/frequently-asked-questions)")

if st.sidebar.button("Start New Chat"):
    st.session_state['responses'] = ["How can I assist you?"]
    st.session_state['requests'] = []
    st.session_state['metadata'] = []
    st.session_state.buffer_memory.clear()

st.title("MeitY Chatbot")

# Load API key
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Initialize LLM and embeddings
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro-latest", google_api_key=GOOGLE_API_KEY)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# Initialize vector store
db = Chroma(
    persist_directory="emb",
    embedding_function=embeddings
)

# Initialize retriever
retriever = RedundantFilterRetriever(
    embeddings=embeddings,
    chroma=db
)

# Session state setup
if 'responses' not in st.session_state:
    st.session_state['responses'] = ["How can I assist you?"]

if 'requests' not in st.session_state:
    st.session_state['requests'] = []

if 'buffer_memory' not in st.session_state:
    st.session_state.buffer_memory = ConversationBufferWindowMemory(k=3, return_messages=True, memory_key="history")

if 'metadata' not in st.session_state:
    st.session_state['metadata'] = []

if "copied" not in st.session_state: 
    st.session_state.copied = []

# Define message templates
system_msg_template = SystemMessagePromptTemplate.from_template(
    template="""Answer the question as relevant as possible using the provided context on MIETY guidelines implementing services, 
    Do not use the internet or external sources. When crafting your response, use as much text as possible from the "response" section of the provided document context, making minimal changes to the original text. 
    Be mindful of case sensitivity differences between the query and context. If the answer is not explicitly contained within the provided context, clearly state 'I don't know' without attempting to fabricate an answer. also dont give use new lines unnecerily giving answers short and sweet.
    Provide a brief response to the user's query first, and then ask for any necessary preferences or clarifications. Do not introduce any additional questions or topics on your own."""
)

human_msg_template = HumanMessagePromptTemplate.from_template(template="{input}")

prompt_template = ChatPromptTemplate.from_messages([
    system_msg_template,
    MessagesPlaceholder(variable_name="history"),
    human_msg_template
])

# Initialize conversation chain
conversation = ConversationChain(
    memory=st.session_state.buffer_memory,
    prompt=prompt_template,
    llm=llm,
    verbose=False
)

# Response container
response_container = st.container()

# Define helper functions
def get_conversation_string():
    conversation_string = ""
    for i in range(len(st.session_state['responses']) - 1):
        conversation_string += "Human: " + st.session_state['requests'][i] + "\n"
        conversation_string += "Bot: " + st.session_state['responses'][i + 1] + "\n"
    return conversation_string

model = GoogleGenerativeAI(model="gemini-pro")

def query_refiner(conversation, query):
    response = model.invoke(f"Given the following user query and conversation log, formulate a question that would be the most relevant to provide the user with an answer from a knowledge base.\n\nCONVERSATION LOG: \n{conversation}\n\nQuery: {query}\n\nRefined Query:")
    return response

def copy_to_clipboard(response):
    pyperclip.copy(response)

# Input text container
query = st.chat_input("How can I guide you!", key="input")
if query:
    conversation_string = get_conversation_string()
    refined_query = query_refiner(conversation_string, query)
    context_docs = db.similarity_search_with_relevance_scores(refined_query, k=2)
    best_source = None
    best_page = None
    best_score = -1
    context_texts = ""

    for doc in context_docs:
        context_texts += doc[0].page_content + "\n" 
        source = doc[0].metadata['source']
        page = doc[0].metadata['page']
        score = doc[0].metadata.get('similarity_score', 0)  

        if score > best_score:
            best_source = source
            best_page = page
            best_score = score

    best_source_text = f"Source: {best_source}, Page: {best_page}"
    response = conversation.predict(input=f"Context:\n {context_texts} \n\n Query:\n {query}")

    st.session_state.requests.append(query)
    st.session_state.responses.append(response)
    st.session_state.metadata.append(best_source_text)

if st.session_state['responses']:
    with st.chat_message("assistant"):
        st.markdown(st.session_state['responses'][0])

# Display the conversation in the chat format
for i in range(len(st.session_state['requests'])):
    # Display user message
    with st.chat_message("user"):
        st.markdown(st.session_state['requests'][i])
    
    # Display assistant response
    with st.chat_message("assistant"):
        st.markdown(st.session_state['responses'][i + 1])
        
        # Add a separator
        st.markdown("---")
        
        # Display metadata and copy button
        metadata_col, button_col = st.columns([3, 1])
        with metadata_col:
            st.markdown(f"{st.session_state['metadata'][i]}")
        with button_col:
            st.button("📋", key=f"copy_button_{i}", on_click=lambda r=st.session_state['responses'][i + 1]: copy_to_clipboard(r), help="Copy response to clipboard")


# Function to export chat to Word document
def export_to_word():
    doc = Document()
    doc.add_heading('Chat History', level=1)
    for i in range(len(st.session_state['responses']) - 1):
        doc.add_heading(f'Q{i + 1}:', level=2)
        doc.add_paragraph(st.session_state['requests'][i])
        doc.add_heading(f'A{i + 1}:', level=2)
        doc.add_paragraph(st.session_state['responses'][i + 1])
        if i < len(st.session_state['metadata']):
            doc.add_paragraph(f'Metadata: {st.session_state["metadata"][i]}')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.sidebar.download_button(
        label="Export Chat as Word Document",
        data=buffer,
        file_name="MeitY_guidelines_chat_history.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Sidebar action
st.sidebar.title("Actions")
export_to_word()
